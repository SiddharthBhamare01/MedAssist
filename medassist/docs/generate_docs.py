"""
Generate CS 595 individual deliverable DOCX files for Siddharth Bhamare.
Produces:
  - Siddharth_CourseExperience.docx
  - Siddharth_ContributionReport.docx
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))


def set_font(run, name="Calibri", size=11, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def set_margins(doc, inches=1.0):
    for section in doc.sections:
        section.top_margin = Inches(inches)
        section.bottom_margin = Inches(inches)
        section.left_margin = Inches(inches)
        section.right_margin = Inches(inches)


def add_heading(doc, text, level=1, size=13, color=(31, 73, 125)):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_font(run, size=size, bold=True, color=color)
    return p


def add_body(doc, text, indent=False, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(space_after)
    if indent:
        p.paragraph_format.left_indent = Inches(0.25)
    run = p.add_run(text)
    set_font(run, size=11)
    return p


def add_concept(doc, number, title, body):
    """Add a numbered concept with bold title inline."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)

    num_run = p.add_run(f"{number}. ")
    set_font(num_run, size=11, bold=True)

    title_run = p.add_run(title)
    set_font(title_run, size=11, bold=True)

    body_run = p.add_run(f"\n{body}")
    set_font(body_run, size=11)
    p.paragraph_format.left_indent = Inches(0.25)
    return p


def add_header_block(doc, title, name, aid, team):
    """Add the document title and info block."""
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.paragraph_format.space_after = Pt(2)
    r = t.add_run("CS 595 — Medical Informatics & AI")
    set_font(r, size=14, bold=True, color=(31, 73, 125))

    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t2.paragraph_format.space_after = Pt(10)
    r2 = t2.add_run(title)
    set_font(r2, size=13, bold=True)

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.paragraph_format.space_after = Pt(14)
    r3 = info.add_run(f"{name}  |  A-ID: {aid}  |  {team}")
    set_font(r3, size=11, color=(89, 89, 89))

    # Horizontal rule
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1F497D")
    pBdr.append(bottom)
    pPr.append(pBdr)


def set_table_borders(table):
    """Add thin borders to all table cells."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "AAAAAA")
        tblBorders.append(el)
    tblPr.append(tblBorders)


# ─────────────────────────────────────────────────────────────
#  DOCUMENT 1 — Course Experience Paper
# ─────────────────────────────────────────────────────────────

def create_course_experience():
    doc = Document()
    set_margins(doc, 1.0)

    add_header_block(
        doc,
        title="Course Experience Paper",
        name="Siddharth Bhamare",
        aid="A20582786",
        team="Team Alpha, Group 2",
    )

    # ── Section 1 ──────────────────────────────────────────────
    add_heading(doc, "Section 1 — Three Concepts Learned in the Overall Course")

    add_concept(
        doc,
        1,
        "HIPAA, HITECH, and Healthcare Privacy Law",
        (
            "Healthcare data is governed by a layered legal framework, not just good intentions. "
            "HIPAA (1996) established the Privacy Rule — defining what constitutes Protected Health "
            "Information (PHI) and who may access it — and the Security Rule, which mandates "
            "administrative, physical, and technical safeguards for electronic PHI. The HITECH Act "
            "(2009) strengthened HIPAA enforcement significantly, introduced the Breach Notification "
            "Rule (covered entities must report breaches to HHS within 60 days), and tied EHR adoption "
            "to Medicare/Medicaid financial incentives, accelerating nationwide digitization. The "
            "21st Century Cures Act (2016) then introduced information-blocking rules — providers "
            "can no longer withhold electronic health data from patients or their designated "
            "applications. Understanding this legal hierarchy was essential to MedAssist: we had to "
            "implement role-based access control, HIPAA-compliant audit trails logging every sensitive "
            "action (user, action, resource, IP, timestamp), and PIN-protected emergency medical ID "
            "cards with brute-force protection. Knowing the why behind these requirements — not just "
            "the how — made every design decision clearer."
        ),
    )

    add_concept(
        doc,
        2,
        "Clinical Decision Support Systems (CDSS) and AI in Medicine",
        (
            "CDSS have existed since the 1970s — MYCIN (1976) used rule-based logic to recommend "
            "antibiotic treatments — but modern AI-driven systems differ in that they use probabilistic "
            "reasoning rather than rigid rule trees. The course covered the FDA's evolving framework "
            "for AI/ML-based Software as a Medical Device (SaMD), which distinguishes between 'locked' "
            "(static, validated) and 'adaptive' (continuously learning) algorithms, and requires "
            "post-market surveillance for the latter. This shaped how we framed MedAssist: the system "
            "provides AI-driven decision support for interpreting blood test results — flagging abnormal "
            "parameters, explaining clinical significance, recommending lifestyle changes, and scheduling "
            "follow-up tests — rather than autonomous medical diagnosis, which would require FDA 510(k) "
            "clearance. Understanding this regulatory boundary kept us from over-claiming what the "
            "system does and grounded the agent output design in appropriate, explainable recommendations."
        ),
    )

    add_concept(
        doc,
        3,
        "Medical Coding Standards: ICD-10, CPT, RxNorm, and FHIR",
        (
            "Interoperability in healthcare depends on everyone using the same vocabulary. ICD-10 "
            "encodes diagnoses and is the basis for insurance billing, epidemiological reporting, and "
            "EHR records. CPT codes encode clinical procedures. RxNorm normalizes drug names — mapping "
            "brand names, generics, and multi-ingredient combinations to a canonical RxCUI identifier. "
            "FHIR (Fast Healthcare Interoperability Resources) is the modern standard for exchanging "
            "EHR data via RESTful APIs, enabling systems built by different vendors to share patient "
            "records without custom integrations. The course explained why these standards exist: "
            "fragmented data across hospital systems directly harms care coordination and epidemiological "
            "research. In MedAssist, medication lookups and tablet recommendations go through RxNorm "
            "for standardized drug identification, and drug adverse event queries use OpenFDA by "
            "NDC/RxCUI. Seeing these standards integrated into a real working product made them "
            "concrete rather than theoretical."
        ),
    )

    # ── Section 2 ──────────────────────────────────────────────
    add_heading(doc, "Section 2 — Three Concepts Learned from Labs and Technical Work")

    add_concept(
        doc,
        1,
        "Government Healthcare APIs and Open Data (OpenFDA and RxNorm)",
        (
            "The labs introduced us to production-grade, freely accessible medical databases maintained "
            "by the US government. OpenFDA exposes drug labels, adverse event reports, and recall "
            "databases at 240 requests per minute with no API key — making it possible to verify "
            "medication safety and look up adverse events without any proprietary subscription. RxNorm "
            "provides drug name normalization, mapping brand names, generics, and multi-ingredient "
            "combinations to a canonical RxCUI identifier so different systems can consistently refer "
            "to the same drug. Hands-on use showed how these APIs can substitute for expensive "
            "proprietary clinical databases, and how to handle their rate limits, inconsistent schemas, "
            "and partial data gracefully. Building real LLM tool-calling agents around these APIs — "
            "rather than just reading about them — was the most impactful part of the lab experience."
        ),
    )

    add_concept(
        doc,
        2,
        "NLP and Vision OCR for Medical Document Processing",
        (
            "The technical sessions highlighted how most clinical information lives in unstructured "
            "form: doctor notes, discharge summaries, and scanned lab reports. The labs covered "
            "techniques for extracting structured data — regex-based parsing, PDF text extraction, "
            "and vision-based OCR. In MedAssist we applied pdf-parse for text-based PDFs and Google "
            "Gemini 1.5 Vision for image or scanned PDFs, extracting 40+ blood parameters (name, "
            "value, unit, normal range, status) using structured JSON output prompts. The challenge "
            "of handling varied lab report formats — different column orders, handwritten values, "
            "multi-page documents — was a direct application of the robustness principles covered "
            "in the labs."
        ),
    )

    add_concept(
        doc,
        3,
        "Geospatial Health Data and Location-Based Services",
        (
            "The labs covered how GIS is used in public health — mapping disease outbreaks, "
            "identifying underserved areas, and routing patients to appropriate facilities. We applied "
            "this using OpenStreetMap's Overpass API to query healthcare facilities (clinics, labs, "
            "hospitals, pharmacies) within a 10 km radius of the patient's location. Implementing "
            "Haversine distance sorting, a 3-mirror API fallback chain with 14-second timeouts, and "
            "1-hour in-memory caching showed how even 'simple' location features require resilient "
            "engineering when the underlying data source is a community-run service rather than a "
            "commercial provider with SLA guarantees."
        ),
    )

    # ── Section 3 ──────────────────────────────────────────────
    add_heading(doc, "Section 3 — Three Concepts Learned from Project Work")

    add_concept(
        doc,
        1,
        "Multi-LLM Ensemble Architecture with Consensus Judging",
        (
            "The most significant technical concept I internalized building MedAssist was that "
            "single-model LLM inference is unreliable for high-stakes medical domains. We built an "
            "ensemble runner that dispatches identical medical prompts to multiple free providers "
            "(SambaNova Llama 3.3 70B, GitHub Models GPT-4o mini, OpenRouter) in parallel using "
            "Promise.allSettled(), then uses OpenAI GPT-4o as a consensus judge to merge outputs "
            "with task-specific strategies. For drug interactions, the system uses the MORE SEVERE "
            "rating when agents disagree. For blood report medical analysis, parallel outputs are "
            "merged conservatively — a parameter is flagged abnormal if any agent raises concern, "
            "rather than requiring full consensus, prioritizing patient safety over false reassurance. "
            "This architecture meaningfully reduces hallucination risk compared to single-call "
            "inference and is the approach I would carry into any future medical AI project."
        ),
    )

    add_concept(
        doc,
        2,
        "Vision AI for Structured Medical Document Extraction",
        (
            "Integrating Gemini 1.5 Vision for blood report OCR taught me how vision-language models "
            "handle structured document understanding differently from conversational chat. The prompt "
            "must describe the expected output schema precisely — 40+ parameter names, value, unit, "
            "normal range, status — and the model must be explicitly instructed to handle missing "
            "values, partial rows, and non-standard units gracefully rather than hallucinating plausible "
            "values. The biggest learning: OCR accuracy for medical documents depends heavily on "
            "prompt engineering, not just model capability. A well-structured schema prompt with "
            "explicit null-handling instructions outperformed generic extraction prompts significantly."
        ),
    )

    add_concept(
        doc,
        3,
        "Evidence-Based Clinical Risk Scoring Algorithms",
        (
            "Implementing Framingham cardiovascular, FINDRISC diabetes, CKD-EPI kidney function, and "
            "Child-Pugh liver scoring required reading primary medical literature and clinical "
            "guidelines — not just technical documentation. These formulas encode decades of "
            "epidemiological research into numerical risk estimates derived from large patient cohorts. "
            "Combining the four scores into a composite 0-100 health index and mapping it to urgency "
            "levels (routine / follow-up soon / urgent / emergency) bridged clinical medicine and "
            "software engineering in a way no other part of the project did. It also reinforced that "
            "medical AI must be grounded in validated clinical science, not just model output."
        ),
    )

    out_path = os.path.join(OUTPUT_DIR, "Siddharth_CourseExperience.docx")
    doc.save(out_path)
    print(f"Saved: {out_path}")


# ─────────────────────────────────────────────────────────────
#  DOCUMENT 2 — Contribution Report
# ─────────────────────────────────────────────────────────────

MEMBERS = [
    {
        "name": "Siddharth Bhamare",
        "rating": "9 / 10",
        "text": (
            "Primary architect and full-stack developer for the project. Designed and implemented the "
            "entire system end-to-end: Express.js backend (API route modules, AI agents, "
            "ensemble runner, service modules), React frontend (15 pages, 13 components, 2 global "
            "contexts), PostgreSQL schema across base schema and migrations, deployment "
            "pipeline (Render + Vercel + Supabase), and both core AI agent systems (blood report "
            "multi-phase analysis and clinical risk scoring). Led all major technical decisions: "
            "multi-provider LLM ensemble with consensus judging, HIPAA audit trail, TOTP 2FA, "
            "Gemini 1.5 Vision OCR for 40+ blood parameters, ElevenLabs TTS narration, Puppeteer "
            "PDF export, and EN/ES internationalization with smart translation caching in PostgreSQL. "
            "Responsible for approximately 90% of the codebase. Rating is 9 rather than 10 "
            "because better time planning in the mid-sprint phase could have allowed more thorough "
            "integration testing and documentation before the final week."
        ),
    },
    {
        "name": "Vaishnav Bhujbhal",
        "rating": "8 / 10",
        "text": (
            "Solid and dependable contributor throughout the project. Provided structured feedback on "
            "the blood report upload and analysis UX flow, which shaped the final page layout and "
            "section ordering of the analysis results. Assisted with manually testing the blood report "
            "upload pipeline across multiple file types (PDF, JPEG, PNG) and flagged two edge cases "
            "in the OCR extraction path that were subsequently fixed. Helped debug a React Router v7 "
            "nested route conflict during integration that was blocking the patient analysis page. "
            "Contributed to the NABC presentation slides — particularly the 'Need' and 'Approach' "
            "sections — and participated actively in all team design discussions. A reliable team "
            "member who contributed meaningfully across both technical testing and presentation work."
        ),
    },
    {
        "name": "Gayatri Gaikwad",
        "rating": "8 / 10",
        "text": (
            "Consistent contributor across documentation, project coordination, and QA. Authored the "
            "initial project definition document and maintained meeting notes that kept the team "
            "aligned on milestones. On the technical side, helped test the patient dashboard and "
            "report history pages — verifying that Recharts trend visualizations rendered correctly "
            "across multiple mock data entries and checking that backdated report timestamps displayed "
            "in the correct month-by-month order. Contributed to the NotebookLM presentation content, "
            "organizing the feature walkthrough section and preparing speaker notes for the demo "
            "portions. Helped prepare the final demo script. A reliable team member who kept the "
            "group on schedule with deliverable checklists and deadline tracking."
        ),
    },
    {
        "name": "Jeevan Singh",
        "rating": "8 / 10",
        "text": (
            "Contributed consistently to testing and presentation throughout the project lifecycle. "
            "Validated the patient dashboard and vitals tracker flows end-to-end and tested the agent "
            "status polling UI, confirming that the 5-phase progress bars (OCR, Medical Analysis, "
            "Lifestyle, Risk Scoring, Follow-Up) updated correctly during live blood report analysis "
            "sessions. Tested the public medical ID card feature for correctness — including PIN "
            "validation, brute-force lockout after 10 attempts, and the emergency contact card display "
            "format. Assisted with scripting and narrating the final demo video walkthrough. Contributed "
            "to the NABC slides and helped finalize the visual layout of the NotebookLM presentation. "
            "A consistent team member who took clear ownership of the QA and presentation deliverables."
        ),
    },
]


def create_contribution_report():
    doc = Document()
    set_margins(doc, 1.0)

    add_header_block(
        doc,
        title="Student Final Project Contribution Report",
        name="Siddharth Bhamare",
        aid="A20582786",
        team="Team Alpha, Group 2",
    )

    # Ratings summary table
    add_heading(doc, "Team Ratings Summary", size=12)

    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for cell, txt in zip(hdr, ["Team Member", "Rating"]):
        cell.paragraphs[0].clear()
        r = cell.paragraphs[0].add_run(txt)
        set_font(r, size=11, bold=True)
        cell.paragraphs[0].paragraph_format.space_before = Pt(3)
        cell.paragraphs[0].paragraph_format.space_after = Pt(3)

    for m in MEMBERS:
        row = table.add_row().cells
        for cell, txt in zip(row, [m["name"], m["rating"]]):
            cell.paragraphs[0].clear()
            r = cell.paragraphs[0].add_run(txt)
            set_font(r, size=11)
            cell.paragraphs[0].paragraph_format.space_before = Pt(3)
            cell.paragraphs[0].paragraph_format.space_after = Pt(3)

    set_table_borders(table)

    # Column widths
    for row in table.rows:
        row.cells[0].width = Inches(3.5)
        row.cells[1].width = Inches(1.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # Individual assessments
    add_heading(doc, "Individual Assessments")

    for m in MEMBERS:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(2)
        name_run = p.add_run(m["name"])
        set_font(name_run, size=12, bold=True, color=(31, 73, 125))
        rating_run = p.add_run(f"  —  {m['rating']}")
        set_font(rating_run, size=12, bold=False, color=(89, 89, 89))

        body = doc.add_paragraph()
        body.paragraph_format.space_before = Pt(2)
        body.paragraph_format.space_after = Pt(8)
        body.paragraph_format.left_indent = Inches(0.2)
        r = body.add_run(m["text"])
        set_font(r, size=11)

    out_path = os.path.join(OUTPUT_DIR, "Siddharth_ContributionReport.docx")
    doc.save(out_path)
    print(f"Saved: {out_path}")


if __name__ == "__main__":
    create_course_experience()
    create_contribution_report()
    print("\nDone. Both DOCX files created in medassist/docs/")
