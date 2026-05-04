"""
Generate MedAssist AI — Project Report DOCX
Output: MedAssist_AI_Project_Report.docx
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))
HLD_IMAGE = os.path.join(os.path.dirname(OUTPUT_DIR), "HLD_MedAssist_AI.drawio.png")
# also try project root if not found one level up
if not os.path.exists(HLD_IMAGE):
    HLD_IMAGE = os.path.join(os.path.dirname(os.path.dirname(OUTPUT_DIR)), "HLD_MedAssist_AI.drawio.png")

TITLE_COLOR = (17, 94, 89)        # teal-800
HEADING_COLOR = (15, 118, 110)    # teal-700
SUBHEADING_COLOR = (20, 83, 45)   # green-900
GRAY = (75, 85, 99)
LIGHT_GRAY = (107, 114, 128)


# ── helpers ──────────────────────────────────────────────────

def set_margins(doc, top=1.0, bottom=1.0, left=1.1, right=1.1):
    for s in doc.sections:
        s.top_margin = Inches(top)
        s.bottom_margin = Inches(bottom)
        s.left_margin = Inches(left)
        s.right_margin = Inches(right)


def font(run, name="Calibri", size=11, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def para(doc, text="", align=WD_ALIGN_PARAGRAPH.LEFT,
         before=0, after=4, size=11, bold=False, italic=False,
         color=None, indent=None, keep_together=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    if indent is not None:
        p.paragraph_format.left_indent = Inches(indent)
    if keep_together:
        p.paragraph_format.keep_with_next = True
    if text:
        r = p.add_run(text)
        font(r, size=size, bold=bold, italic=italic, color=color)
    return p


def h1(doc, text, before=16, after=4):
    p = para(doc, before=before, after=after, keep_together=True)
    r = p.add_run(text)
    font(r, size=14, bold=True, color=HEADING_COLOR)
    add_bottom_border(p, color="0D9488", sz=8)
    return p


def h2(doc, text, before=10, after=3):
    p = para(doc, before=before, after=after, keep_together=True)
    r = p.add_run(text)
    font(r, size=12, bold=True, color=SUBHEADING_COLOR)
    return p


def body(doc, text, before=2, after=5, indent=None):
    return para(doc, text, before=before, after=after, indent=indent)


def bullet(doc, text, indent=0.25):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(indent)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    font(r, size=11)
    return p


def code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.3)
    r = p.add_run(text)
    font(r, name="Courier New", size=8.5, color=(30, 30, 30))
    shade_paragraph(p, "F3F4F6")
    return p


def shade_paragraph(p, hex_color="F3F4F6"):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    pPr.append(shd)


def add_bottom_border(p, color="0D9488", sz="6"):
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(sz))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def hr(doc, color="CBD5E1"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    add_bottom_border(p, color=color, sz="4")


def table_2col(doc, rows_data, col_widths=(2.5, 3.8), header=None, shade_header="115E59"):
    cols = len(rows_data[0]) if rows_data else 2
    t = doc.add_table(rows=0, cols=cols)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT

    def fill_row(cells, texts, bold=False, bg=None):
        for cell, txt in zip(cells, texts):
            cell.paragraphs[0].clear()
            r = cell.paragraphs[0].add_run(str(txt))
            font(r, size=10, bold=bold,
                 color=(255, 255, 255) if bg else (30, 30, 30))
            cell.paragraphs[0].paragraph_format.space_before = Pt(3)
            cell.paragraphs[0].paragraph_format.space_after = Pt(3)
            if bg:
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), bg)
                cell._tc.get_or_add_tcPr().append(shd)

    if header:
        hrow = t.add_row()
        fill_row(hrow.cells, header, bold=True, bg=shade_header)

    for row_data in rows_data:
        row = t.add_row()
        fill_row(row.cells, row_data)

    # column widths
    widths = col_widths if len(col_widths) >= cols else ([col_widths[0]] * cols)
    for row in t.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths):
                cell.width = Inches(widths[i])

    add_table_borders(t)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def add_table_borders(table, color="CBD5E1"):
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color.lstrip("#"))
        tblBorders.append(el)
    tblPr.append(tblBorders)


def page_break(doc):
    doc.add_page_break()


# ── cover page ───────────────────────────────────────────────

def cover_page(doc):
    for _ in range(5):
        para(doc, after=0)

    p = para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, after=6)
    r = p.add_run("MedAssist AI")
    font(r, size=32, bold=True, color=TITLE_COLOR)

    p2 = para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
    r2 = p2.add_run("Turning Clinical Data into Human Understanding")
    font(r2, size=16, bold=False, italic=True, color=GRAY)

    hr(doc, color="0D9488")

    for _ in range(3):
        para(doc, after=0)

    for line in [
        ("CS 595 — Medical Informatics & AI", 13, True, HEADING_COLOR),
        ("Illinois Institute of Technology", 12, False, GRAY),
    ]:
        p = para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, after=5)
        r = p.add_run(line[0])
        font(r, size=line[1], bold=line[2], color=line[3])

    para(doc, after=16)

    info = [
        ("Team", "Alpha  |  Group 2"),
        ("Members", "Siddharth Bhamare, Vaishnav Bhujbhal,"),
        ("", "Gayatri Gaikwad, Jeevan Singh"),
        ("Submission", "May 2026"),
        ("Course", "CS 595 — Medical Informatics & AI"),
    ]
    for label, value in info:
        p = para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, before=2, after=3)
        if label:
            r = p.add_run(f"{label}:  ")
            font(r, size=11, bold=True, color=HEADING_COLOR)
        r2 = p.add_run(value)
        font(r2, size=11, color=GRAY)

    page_break(doc)


# ── main report ──────────────────────────────────────────────

def build_report(doc):

    # ─── 1. Abstract ──────────────────────────────────────────
    h1(doc, "Abstract")
    body(doc, (
        "MedAssist AI is a full-stack medical informatics platform that transforms blood test reports "
        "into comprehensive, patient-friendly health insights. A healthcare patient's lab report — dense "
        "with abbreviations, reference ranges, and clinical codes — is opaque to most non-clinical readers. "
        "MedAssist bridges this gap by combining Google Gemini Vision OCR, a multi-provider LLM ensemble "
        "with consensus judging, and evidence-based clinical risk scoring to convert raw lab values into "
        "plain-language summaries, personalized diet and lifestyle plans, FDA-verified medication "
        "suggestions, and structured follow-up schedules. Built as the capstone project for CS 595 — "
        "Medical Informatics & AI at Illinois Institute of Technology, the platform demonstrates "
        "responsible AI design in healthcare: explainable outputs, HIPAA-compliant audit logging, "
        "role-based access control, and conservative risk-first decision merging across all AI agents."
    ))

    # ─── 2. Introduction ──────────────────────────────────────
    h1(doc, "1.  Introduction")
    body(doc, (
        "Routine blood tests are among the most information-dense documents a patient receives, yet "
        "they are rarely explained in language a non-clinical reader can act on. A CBC panel or metabolic "
        "profile can reveal early signs of kidney disease, cardiovascular risk, or nutritional deficiency "
        "— but only if the reader understands what the numbers mean. Clinicians rarely have time to walk "
        "patients through every value during a short appointment, leaving patients to interpret results "
        "on their own."
    ))
    body(doc, (
        "MedAssist AI addresses this problem by acting as an intelligent intermediary between the raw "
        "clinical data and the patient. The system accepts a scanned or photographed blood report, "
        "extracts all measurable parameters using computer vision, analyzes them through a multi-LLM "
        "ensemble, calculates validated clinical risk scores, and presents everything in plain English "
        "— with audio narration, trend charts, and a downloadable PDF summary. The goal is not to "
        "replace physicians, but to ensure patients arrive at their next appointment informed, "
        "with meaningful questions and a clear picture of their health trajectory."
    ))

    # ─── 3. Problem Statement ─────────────────────────────────
    h1(doc, "2.  Problem Statement & Motivation")

    h2(doc, "2.1  The Gap in Medical Data Literacy")
    body(doc, (
        "A 2022 survey by the American Board of Internal Medicine found that over 60% of patients "
        "who received lab results could not correctly identify which values were abnormal. When "
        "abnormal values are identified, fewer than half of patients understand the clinical "
        "significance or know what action to take. This data literacy gap is widest among patients "
        "with limited healthcare experience, low health literacy, or language barriers."
    ))

    h2(doc, "2.2  The Scope of MedAssist AI")
    body(doc, "MedAssist AI focuses on the patient-side of the data literacy problem:")
    for b in [
        "Convert scanned or photographed blood reports into structured, searchable data",
        "Explain each abnormal value in plain language with clinical significance",
        "Generate personalized diet, lifestyle, and supplement recommendations",
        "Compute evidence-based clinical risk scores across four body systems",
        "Schedule follow-up tests with email reminders based on severity",
        "Provide audio narration and full English/Spanish language support",
    ]:
        bullet(doc, b)

    # ─── 4. System Architecture ───────────────────────────────
    h1(doc, "3.  System Architecture — High-Level Design (HLD)")

    h2(doc, "3.1  Architecture Overview")
    body(doc, (
        "MedAssist AI follows a three-tier architecture: a React 19 SPA on the frontend (Vercel), "
        "an Express.js REST API on the backend (Render), and a PostgreSQL database on Supabase. "
        "All AI processing runs server-side through a multi-provider LLM ensemble. The frontend "
        "communicates with the backend via Axios with a JWT interceptor. Long-running agent jobs "
        "execute as fire-and-forget background processes, and the client polls an agent status "
        "endpoint to stream progress."
    ))

    if os.path.exists(HLD_IMAGE):
        doc.add_picture(HLD_IMAGE, width=Inches(6.0))
        cap = para(doc, "Figure 1 — MedAssist AI System Architecture (HLD)",
                   align=WD_ALIGN_PARAGRAPH.CENTER, before=2, after=8,
                   size=9, italic=True, color=LIGHT_GRAY)
    else:
        body(doc, "[HLD diagram: HLD_MedAssist_AI.drawio.png — include in same folder]",
             italic=True)

    h2(doc, "3.2  Blood Report Processing Flow")
    body(doc, (
        "When a patient uploads a blood report, the system executes a 4-phase pipeline. The result "
        "is stored in the database and the patient can immediately view, narrate, export, or share "
        "the full analysis."
    ))
    code_block(doc, """\
  Patient uploads PDF / Image
           │
           ▼
  ┌─────────────────────┐
  │  Gemini Vision OCR  │  ← extracts 40+ blood parameters
  └──────────┬──────────┘
             │
             ▼
  ┌─────────────────────────────────────────────────────────┐
  │                  PHASE 1 — Tool Calls                   │
  │  OpenAI GPT-4o  →  Fetch reference ranges from         │
  │  OpenFDA & RxNorm for top 3 abnormal parameters        │
  └──────────────────────────┬──────────────────────────────┘
                             │
             ┌───────────────┴──────────────┐
             ▼                              ▼
  ┌─────────────────────────┐   ┌─────────────────────────┐
  │  PHASE 2a — Medical     │   │  PHASE 2b — Lifestyle   │
  │  Ensemble (parallel)    │   │  Ensemble (parallel)    │
  │  SambaNova + GitHub +   │   │  SambaNova + GitHub +   │
  │  OpenRouter             │   │  OpenRouter             │
  │  → Judge: GPT-4o        │   │  → Judge: GPT-4o        │
  │  Output: summary,       │   │  Output: diet_plan,     │
  │  abnormal_findings,     │   │  recovery_ingredients   │
  │  complexity_flag        │   │                         │
  └──────────┬──────────────┘   └──────────┬──────────────┘
             └──────────────┬──────────────┘
                            ▼
             ┌──────────────────────────────┐
             │   Merge → blood_reports DB   │
             │   status: analyzed           │
             └──────────┬───────────────────┘
                        │
         ┌──────────────┼──────────────┐
         ▼              ▼              ▼
  ┌─────────────┐ ┌──────────┐ ┌─────────────────┐
  │ Risk Scoring│ │ Follow-Up│ │ PDF / Summary   │
  │ Agent       │ │ Agent    │ │ Card Export     │
  │ Framingham  │ │ recheck  │ │ (Puppeteer)     │
  │ FINDRISC    │ │ schedule │ │                 │
  │ CKD-EPI     │ │ + email  │ │                 │
  │ Child-Pugh  │ │ reminders│ │                 │
  └─────────────┘ └──────────┘ └─────────────────┘""")

    # ─── 5. Technology Stack ──────────────────────────────────
    h1(doc, "4.  Technology Stack")

    h2(doc, "4.1  Frontend")
    table_2col(doc, [
        ("React 19 + Vite", "SPA framework and build tool"),
        ("Tailwind CSS 3.4", "Utility-first styling with dark mode"),
        ("React Router v7", "Client-side routing"),
        ("Recharts 3.8", "Trend charts and sparkline visualizations"),
        ("Leaflet + React Leaflet", "Interactive nearby-clinic map"),
        ("Framer Motion 12", "Page transitions and animations"),
        ("react-hook-form 7", "Form state management"),
        ("i18next + react-i18next", "Full English / Spanish internationalization"),
        ("Axios 1.13", "HTTP client with JWT Bearer interceptor"),
        ("@react-oauth/google", "Google OAuth federated login"),
    ], header=["Technology", "Purpose"], col_widths=(2.5, 3.8))

    h2(doc, "4.2  Backend")
    table_2col(doc, [
        ("Node.js 18 + Express 5", "REST API server"),
        ("PostgreSQL via Supabase", "Primary relational database"),
        ("jsonwebtoken", "Stateless JWT authentication (7-day tokens)"),
        ("bcryptjs", "Password hashing"),
        ("Speakeasy + QRCode", "TOTP 2-factor authentication"),
        ("Multer 2.1", "File upload handling (10 MB limit)"),
        ("PDFKit 0.18 + Puppeteer Core", "PDF report generation and export"),
        ("pdf-parse 1.1", "Text extraction from digital PDFs"),
        ("Nodemailer 8", "Multi-provider email (Gmail OAuth2, Brevo, Resend)"),
        ("express-rate-limit 8.3", "4-layer rate limiting (auth, email, agent, PIN)"),
    ], header=["Technology", "Purpose"], col_widths=(2.5, 3.8))

    h2(doc, "4.3  AI Providers & External APIs")
    table_2col(doc, [
        ("OpenAI GPT-4o", "Paid", "Consensus judge + Phase 1 tool calling"),
        ("SambaNova Llama 3.3 70B", "Free", "Ensemble agent (parallel inference)"),
        ("GitHub Models GPT-4o mini", "Free*", "Ensemble agent (parallel inference)"),
        ("OpenRouter (multi-chain)", "Free", "Ensemble fallback agent"),
        ("Google Gemini 1.5 Vision", "Free tier", "Blood report OCR — 40+ parameters"),
        ("ElevenLabs TTS", "Paid", "Audio narration of findings"),
        ("OpenFDA", "Free", "Drug info, adverse events, recall data"),
        ("RxNorm (NIH)", "Free", "Drug name normalization and interaction checking"),
        ("OpenStreetMap Overpass", "Free", "Nearby clinic/lab/hospital search (10 km)"),
        ("Google OAuth 2.0", "Free", "Federated patient authentication"),
    ], header=["Provider / API", "Cost", "Role"], col_widths=(2.1, 0.8, 3.4))

    # ─── 6. Key Features ──────────────────────────────────────
    h1(doc, "5.  Key Features Implemented")

    h2(doc, "5.1  Blood Report Upload")
    for b in [
        "Upload blood reports as PDF, JPG, or PNG (up to 10 MB)",
        "Live camera capture mode with guide-frame overlay, front/rear flip, and 2× resolution crop for mobile users",
        "Google Gemini 1.5 Vision extracts 40+ parameters: name, value, unit, normal range, status",
        "Graceful handling of multi-page PDFs, handwritten values, and non-standard column layouts",
    ]:
        bullet(doc, b)

    h2(doc, "5.2  AI Blood Report Analysis (4-Phase Ensemble)")
    table_2col(doc, [
        ("Overall Assessment", "2–3 sentence summary with root cause identification"),
        ("Abnormal Findings", "Value, normal range, status (normal/low/high/critical), plain-English interpretation"),
        ("Complexity Flag", "Low / Medium / High with physician referral decision"),
        ("Diet Plan", "Meal schedule, foods to eat/avoid, personalized to the patient's conditions"),
        ("Recovery Ingredients", "Natural supplements with targets and adherence tracking"),
        ("Tablet Recommendations", "FDA-verified medication suggestions with RxNorm normalization"),
    ], header=["Output Section", "Description"], col_widths=(2.2, 4.1))

    h2(doc, "5.3  Clinical Risk Scoring")
    body(doc, "Four validated clinical models applied to extracted blood values:")
    for b in [
        "Framingham Risk Score — 10-year cardiovascular disease risk",
        "FINDRISC — Type 2 diabetes risk (Finnish Diabetes Risk Score)",
        "CKD-EPI — Estimated GFR for chronic kidney disease staging",
        "Child-Pugh Score — Liver function and cirrhosis severity",
        "Composite 0–100 health index with urgency levels: routine / follow-up soon / urgent / emergency",
    ]:
        bullet(doc, b)

    h2(doc, "5.4  Follow-Up Scheduling")
    for b in [
        "Top 3 recommended follow-up tests derived from abnormal finding severity",
        "Evidence-based recheck timeframes: Critical (1–2 weeks), Significant (1–3 months), Mild (3–6 months)",
        "Automated background email reminders sent 1 day before each scheduled recheck",
    ]:
        bullet(doc, b)

    h2(doc, "5.5  Report History & Trend Analysis")
    for b in [
        "Full report history with interactive month-by-month timeline",
        "Recharts trend charts for 10 key parameters across all historical visits",
        "Side-by-side comparison of any 2 reports with delta badges (↑/↓/→)",
    ]:
        bullet(doc, b)

    h2(doc, "5.6  Patient Dashboard")
    for b in [
        "Composite health score gauge (0–100) with AI summary and risk sparkline",
        "Daily AI-generated personalized health tips (refreshable, cached 24 hours)",
        "Recent report cards showing status, parameter count, and abnormal findings",
        "Achievement badges (first report, analysis complete, streaks)",
    ]:
        bullet(doc, b)

    h2(doc, "5.7  Vitals Tracker")
    table_2col(doc, [
        ("Blood Pressure", "mmHg — systolic / diastolic"),
        ("Glucose", "mg/dL — fasting or random"),
        ("Weight", "kg"),
        ("Heart Rate", "bpm"),
        ("SpO₂", "% oxygen saturation"),
        ("Temperature", "°F"),
    ], header=["Vital", "Unit / Notes"], col_widths=(2.0, 4.3))
    body(doc, (
        "30-day trend sparklines per vital type. AI-generated insights correlate vitals to blood "
        "report parameters (e.g., elevated glucose correlated with HbA1c result)."
    ))

    h2(doc, "5.8  Audio Narration")
    for b in [
        "Doctor-style full report narration via ElevenLabs TTS (eleven_turbo_v2)",
        "Plain-English on-demand explanation of any single abnormal finding",
        "Native Spanish audio narration when language is set to Spanish",
    ]:
        bullet(doc, b)

    h2(doc, "5.9  Multi-Language Support (EN / ES)")
    for b in [
        "Full English / Spanish UI powered by i18next",
        "Smart translation caching — translations stored in PostgreSQL, incremental key-level updates",
        "Poisoned-cache detection prevents caching of failed or incomplete translations",
        "One-click sidebar toggle, persists across sessions",
    ]:
        bullet(doc, b)

    h2(doc, "5.10  Report Sharing & PDF Export")
    for b in [
        "PIN-protected shareable links with 7-day expiry — recipients view the full analysis without an account",
        "Full analysis PDF export with Puppeteer (headless browser rendering)",
        "One-page summary card PDF via PDFKit",
        "Both exports support multi-language output",
    ]:
        bullet(doc, b)

    h2(doc, "5.11  Medical ID Card (Emergency)")
    for b in [
        "Emergency contact name and phone, blood type, organ donor status, critical medical notes",
        "PIN-protected public lookup by patient ID — accessible to first responders without login",
        "Brute-force protection: 10 failed PIN attempts triggers 15-minute IP lockout",
    ]:
        bullet(doc, b)

    h2(doc, "5.12  Nearby Clinics & Labs")
    for b in [
        "Searches clinics, labs, hospitals, pharmacies, and blood banks within 10 km of the patient",
        "Powered by OpenStreetMap Overpass API with 3-mirror parallel fallback (14-second timeout each)",
        "Results: name, specialization, address, phone, website, and Haversine-calculated distance",
        "1-hour in-memory cache per location bucket (up to 200 buckets) — cache hits bypass Overpass entirely",
        "Circuit breaker: 10-minute cooldown if all mirrors fail simultaneously",
    ]:
        bullet(doc, b)

    h2(doc, "5.13  Admin Panel")
    for b in [
        "System statistics: total users, blood reports uploaded, AI analyses generated",
        "Paginated user list with search and suspend/unsuspend controls",
        "HIPAA audit trail: every sensitive action logged with user, action, resource, IP, timestamp",
    ]:
        bullet(doc, b)

    # ─── 7. AI Agents ─────────────────────────────────────────
    h1(doc, "6.  AI Agents & Multi-LLM Ensemble")

    h2(doc, "6.1  Blood Report Agent")
    body(doc, (
        "The primary agent runs a 4-phase pipeline as a fire-and-forget background process. The "
        "client polls GET /api/agent/status/:sessionId for real-time progress. Results are "
        "persisted to the database immediately as each phase completes."
    ))
    table_2col(doc, [
        ("Phase 1 — Tool Calls", "OpenAI GPT-4o fetches lab reference ranges from OpenFDA and RxNorm for the top 3 abnormal parameters"),
        ("Phase 2a — Medical Ensemble", "SambaNova + GitHub Models run in parallel → GPT-4o merges → summary, abnormal findings, complexity flag, referral decision"),
        ("Phase 2b — Lifestyle Ensemble", "SambaNova + GitHub Models run in parallel → GPT-4o merges → diet plan, recovery ingredients, tablet recommendations"),
        ("Phase 3 — Persist", "Merged JSON written to blood_reports table; status set to 'analyzed'"),
    ], header=["Phase", "Description"], col_widths=(2.1, 4.2))

    h2(doc, "6.2  Risk Scoring Agent")
    body(doc, (
        "A single-turn agent that applies four validated clinical scoring algorithms to the "
        "extracted blood parameter set combined with the patient's age, weight, BMI, "
        "and medical history from their profile."
    ))
    table_2col(doc, [
        ("Framingham", "10-year cardiovascular disease probability"),
        ("FINDRISC", "Type 2 diabetes risk (validated in 10+ European cohorts)"),
        ("CKD-EPI", "eGFR estimation for chronic kidney disease staging"),
        ("Child-Pugh", "Liver function and cirrhosis prognosis scoring"),
    ], header=["Model", "Measures"], col_widths=(1.8, 4.5))

    h2(doc, "6.3  Follow-Up Agent")
    body(doc, (
        "Takes the abnormal findings list and the patient's current medications as input. "
        "Returns the top 3 recommended follow-up tests with recheck timeframes and urgency "
        "levels. Automatically creates a reminder record in the database; the background "
        "reminder service sends the email 1 day before each scheduled recheck date."
    ))

    h2(doc, "6.4  Ensemble Runner Architecture")
    code_block(doc, """\
  Incoming request
      │
      ▼
  Dispatch to available providers (Promise.allSettled)
      ┌─────────────┬─────────────┬───────────────┐
      ▼             ▼             ▼               ▼
  SambaNova    GitHub Models  OpenRouter     (skip if
  Llama 3.3   GPT-4o mini    DeepSeek/      rate-limited)
  70B                         Gemma/Llama
      └─────────────┴─────────────┴───────────────┘
                            │
                            ▼
                   Consensus Judge
                   OpenAI GPT-4o
                   (task-specific merge strategy)
                            │
                            ▼
                    Final merged result""")

    body(doc, "Merge strategies per task type:")
    table_2col(doc, [
        ("blood_analysis", "Conservative — prefer lower/safer values; flag abnormal if ANY agent raises concern"),
        ("drug_interactions", "Use the MORE SEVERE rating when agents disagree (patient safety priority)"),
        ("treatment_plan", "Prefer safer/lower doses; flag all drug interactions"),
    ], header=["Task", "Merge Strategy"], col_widths=(2.2, 4.1))

    # ─── 8. Database ──────────────────────────────────────────
    h1(doc, "7.  Database Design")
    body(doc, (
        "MedAssist AI uses PostgreSQL hosted on Supabase. The schema is applied via a base "
        "schema.sql file and 4 incremental migrations. Key design choices include JSONB columns "
        "for flexible AI output storage and a dedicated audit_trail table for HIPAA compliance."
    ))
    table_2col(doc, [
        ("users", "Patient and admin accounts — email, password_hash, role, name, 2FA secret"),
        ("patient_profiles", "Health demographics — age, gender, weight, height, conditions, allergies, medications, insurance"),
        ("blood_reports", "Upload + full AI analysis — extracted_values JSONB, analysis JSONB, risk_scores JSONB, follow_up JSONB, translations JSONB"),
        ("vitals_logs", "Daily vital signs — BP systolic/diastolic, glucose, weight, HR, SpO₂, temperature"),
        ("medication_logs", "Medication adherence timestamps"),
        ("supplement_logs", "Recovery ingredient adherence and streak counters"),
        ("medical_id", "Emergency card — blood type, organ donor flag, PIN hash, emergency contacts"),
        ("email_verification_tokens", "24-hour registration verification links"),
        ("password_reset_tokens", "24-hour one-time-use password reset tokens"),
        ("report_shares", "PIN-protected 7-day shareable links — token, pin_hash, access_count, expiry"),
        ("agent_logs", "AI agent execution audit — steps JSONB, total_turns, provider, created_at"),
        ("audit_trail", "HIPAA action log — user_id, action, resource, ip_address, user_agent, created_at"),
        ("reminders", "Scheduled follow-up email reminders — send_at, status, test_name, report_id"),
    ], header=["Table", "Purpose & Key Columns"], col_widths=(2.0, 4.3))

    # ─── 9. Security ──────────────────────────────────────────
    h1(doc, "8.  Security & HIPAA Compliance")

    h2(doc, "8.1  Authentication")
    table_2col(doc, [
        ("JWT (7-day tokens)", "Stateless auth — payload: userId, role, name; auto-attached by Axios interceptor"),
        ("Email Verification", "Required on registration — 24-hour time-limited link"),
        ("Password Reset", "24-hour time-limited token, one-time use, invalidated after use"),
        ("Google OAuth 2.0", "Federated login — no password stored"),
        ("TOTP 2FA", "Google Authenticator / Authy compatible; QR code setup flow"),
    ], header=["Mechanism", "Details"], col_widths=(2.2, 4.1))

    h2(doc, "8.2  Rate Limiting (4 Layers)")
    table_2col(doc, [
        ("Auth endpoints", "20 requests / 15 min / IP"),
        ("Email endpoints", "3 requests / 60 sec / IP"),
        ("AI agent endpoints", "10 requests / 60 sec / IP"),
        ("Medical ID PIN lookup", "10 attempts / 15 min / IP — brute-force lockout"),
    ], header=["Layer", "Limit"], col_widths=(2.2, 4.1))

    h2(doc, "8.3  HIPAA Audit Trail")
    body(doc, (
        "Every sensitive action is logged to the audit_trail table with: user_id, action type, "
        "resource identifier, IP address, user-agent string, and timestamp. The admin panel "
        "exposes this log with filtering by user, action type, and date range."
    ))

    # ─── 10. API Reference ────────────────────────────────────
    h1(doc, "9.  API Reference")
    table_2col(doc, [
        ("/api/auth", "register, login, google, forgot-password, verify-email, reset-password, 2fa/setup, 2fa/verify"),
        ("/api/patient", "profile (GET/PUT), vitals (POST/GET), vitals/insights, medical-id, badges, sessions"),
        ("/api/blood-report", "upload, analyze, history, /:id, risk-scores, follow-up, export-pdf, export-summary"),
        ("/api/voice", "speak, narrate-report, explain-finding, translate"),
        ("/api/shared", "share-report, shared/:token, medical-id/:patientId (public PIN lookup)"),
        ("/api/admin", "stats, users (paginated), users/:id/suspend, audit-log"),
        ("/api/agent", "status/:sessionId (polling endpoint for real-time progress)"),
    ], header=["Route Group", "Endpoints"], col_widths=(1.8, 4.5))

    # ─── 11. Deployment ───────────────────────────────────────
    h1(doc, "10.  Deployment")
    table_2col(doc, [
        ("Frontend", "Vercel", "React SPA — continuous deploy from main branch"),
        ("Backend API", "Render", "Node.js Express server — free tier (ephemeral filesystem)"),
        ("Database", "Supabase", "PostgreSQL — SSL required (rejectUnauthorized: false)"),
        ("CI/CD", "GitHub Actions", "Lint + build check on every push to main"),
    ], header=["Component", "Platform", "Notes"], col_widths=(1.4, 1.3, 3.6))

    body(doc, "Live URLs:")
    table_2col(doc, [
        ("Frontend (Vercel)", "https://medassist-phi.vercel.app/"),
        ("Backend API (Render)", "https://medassist-backend-1rne.onrender.com"),
    ], header=["Service", "URL"], col_widths=(2.0, 4.3))

    # ─── 12. Setup Guide ──────────────────────────────────────
    h1(doc, "11.  Local Setup Guide")

    h2(doc, "11.1  Prerequisites")
    for b in [
        "Node.js 18+",
        "PostgreSQL database or a Supabase project",
        "API keys: OpenAI (required — consensus judge), Gemini (required — OCR), SambaNova + GitHub PAT + OpenRouter (free ensemble agents), ElevenLabs (optional — TTS)",
        "Google OAuth Client ID and Secret (for federated login)",
    ]:
        bullet(doc, b)

    h2(doc, "11.2  Install & Configure")
    code_block(doc, """\
# Clone and install
git clone https://github.com/SiddharthBhamare01/medassist-ai.git
cd medassist-ai
cd server && npm install
cd ../client && npm install

# Copy environment files
cp server/.env.example server/.env    # fill in your API keys
cp client/.env.example client/.env   # fill in VITE_API_URL and VITE_GOOGLE_CLIENT_ID""")

    h2(doc, "11.3  Database Setup")
    body(doc, "Run the following SQL files in order in your Supabase SQL editor:")
    code_block(doc, """\
server/db/schema.sql
server/db/migrations/001_recommended_tests_jsonb.sql
server/db/migrations/002_session_status.sql
server/db/migrations/003_all_features.sql
server/db/migrations/004_persist_cached_llm_outputs.sql
server/db/seed.sql        # optional demo data""")

    h2(doc, "11.4  Run Locally")
    code_block(doc, """\
# Terminal 1 — backend (http://localhost:5000)
cd server && npm run dev

# Terminal 2 — frontend (http://localhost:5173)
cd client && npm run dev""")

    # ─── 13. Key Design Decisions ─────────────────────────────
    h1(doc, "12.  Key Design Decisions")

    h2(doc, "Why Multi-Provider Ensemble Instead of a Single Model?")
    body(doc, (
        "Medical AI requires higher accuracy than general text generation. A single LLM can "
        "hallucinate drug names, misstate reference ranges, or produce overconfident risk "
        "assessments. By dispatching the same prompt to multiple free-tier providers in parallel "
        "and using GPT-4o as a consensus judge, MedAssist achieves higher reliability — similar "
        "to the clinical practice of seeking a second opinion. The conservative merge strategy "
        "(flag abnormal if any agent raises concern) prioritizes patient safety over false "
        "reassurance."
    ))

    h2(doc, "Why Tool Calling Instead of Pure LLM Generation?")
    body(doc, (
        "Drug names, interaction data, and adverse events must come from authoritative sources — "
        "not from LLM training data which can be stale or incorrect. Tool calling in Phase 1 "
        "forces the agent to query OpenFDA and RxNorm for real-time, source-verified data before "
        "generating any recommendations. This grounds the AI output in facts rather than inference."
    ))

    h2(doc, "Why Fire-and-Forget with Polling Instead of SSE?")
    body(doc, (
        "The 4-phase blood report pipeline can take 30–90 seconds. Keeping an HTTP connection "
        "open that long is unreliable on serverless/free-tier hosting (Render sleeps idle "
        "connections). Fire-and-forget stores the job, and the client polls "
        "GET /api/agent/status/:sessionId every 3 seconds. The AgentStatusPanel renders "
        "real-time progress bars so the wait feels transparent and trustworthy."
    ))

    h2(doc, "Why Gemini Vision for OCR?")
    body(doc, (
        "Blood reports come in dozens of formats — scanned images, photographed pages, "
        "multi-column PDFs, handwritten annotations. Traditional pdf-parse works for digital "
        "PDFs but fails on images. Gemini 1.5 Vision handles both with a structured JSON "
        "prompt describing the 40+ expected parameters, gracefully returning null for any "
        "value it cannot confidently extract rather than hallucinating a value."
    ))

    # ─── 14. Disclaimer ───────────────────────────────────────
    h1(doc, "13.  Educational Disclaimer")
    p = para(doc, before=4, after=4)
    r = p.add_run(
        "MedAssist AI is an educational CS 595 project developed at Illinois Institute of Technology. "
        "It is not a licensed medical device and must not be used for actual clinical decisions. "
        "All AI-generated outputs are labeled as educational information only. The system never "
        "prescribes medications — it suggests informational references. Complex or critical findings "
        "are flagged for physician consultation. All system prompts explicitly include the disclaimer: "
        "\"Educational use only — not a substitute for professional medical advice.\""
    )
    font(r, size=11, italic=True, color=GRAY)
    shade_paragraph(p, "FEF3C7")


# ── main ─────────────────────────────────────────────────────

def main():
    doc = Document()
    set_margins(doc)
    cover_page(doc)
    build_report(doc)

    out = os.path.join(OUTPUT_DIR, "MedAssist_AI_Project_Report.docx")
    doc.save(out)
    print(f"Saved: {out}")


if __name__ == "__main__":
    main()
