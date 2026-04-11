"""
Generate MedAssist AI system flowchart and insert into the proposal docx.
Uses only Pillow (PIL) + python-docx — no matplotlib needed.
"""

from PIL import Image, ImageDraw, ImageFont
import math, os

# ── Canvas ─────────────────────────────────────────────────────────────────
W, H = 1540, 1160
img = Image.new("RGB", (W, H), "#FAFAFA")
draw = ImageDraw.Draw(img)

# ── Palette ────────────────────────────────────────────────────────────────
DARK_BLUE   = "#1F3864"
MID_BLUE    = "#2E74B5"
LIGHT_BLUE  = "#D6E4F7"
AGENT_BG    = "#1F3864"
AGENT_FG    = "#FFFFFF"
RESULT_BG   = "#1F6B8E"
RESULT_FG   = "#FFFFFF"
PROCESS_BG  = "#BDD7EE"
PROCESS_FG  = "#1F3864"
TERMINAL_BG = "#1F3864"
TERMINAL_FG = "#FFFFFF"
DECISION_BG = "#FFF2CC"
DECISION_BD = "#BF8F00"
DECISION_FG = "#7F6000"
API_BG      = "#375623"
API_FG      = "#FFFFFF"
ARROW_COL   = "#404040"
DIVIDER_COL = "#CCCCCC"

# ── Fonts ──────────────────────────────────────────────────────────────────
BASE = "C:/Windows/Fonts/"
def lf(name, size):
    for candidate in [BASE + name, BASE + name.lower()]:
        if os.path.exists(candidate):
            return ImageFont.truetype(candidate, size)
    return ImageFont.load_default()

fn   = lf("Calibri.ttf",  18)
fb   = lf("Calibrib.ttf", 18)
fsm  = lf("Calibri.ttf",  15)
fbsm = lf("Calibrib.ttf", 15)
flg  = lf("Calibrib.ttf", 22)
fxlg = lf("Calibrib.ttf", 26)

# ── Drawing primitives ─────────────────────────────────────────────────────
def text_center(bbox, lines, font, color):
    """Draw multi-line text centered in bbox (x1,y1,x2,y2)."""
    x1, y1, x2, y2 = bbox
    lh = font.size + 3
    total = len(lines) * lh - 3
    cy = (y1 + y2) / 2 - total / 2
    for line in lines:
        bb = draw.textbbox((0, 0), line, font=font)
        tw = bb[2] - bb[0]
        draw.text(((x1+x2)/2 - tw/2, cy), line, fill=color, font=font)
        cy += lh

def process_box(cx, y, w, h, lines, bg=PROCESS_BG, fg=PROCESS_FG, font=None, radius=6):
    x1, y1, x2, y2 = cx - w//2, y, cx + w//2, y + h
    draw.rounded_rectangle([x1,y1,x2,y2], radius=radius, fill=bg, outline=MID_BLUE, width=2)
    text_center((x1,y1,x2,y2), lines, font or fb, fg)
    return y2

def terminal(cx, y, w, h, lines, bg=TERMINAL_BG, fg=TERMINAL_FG):
    x1, y1, x2, y2 = cx - w//2, y, cx + w//2, y + h
    draw.rounded_rectangle([x1,y1,x2,y2], radius=h//2, fill=bg, outline=DARK_BLUE, width=2)
    text_center((x1,y1,x2,y2), lines, fb, fg)
    return y2

def agent_box(cx, y, w, h, lines):
    x1, y1, x2, y2 = cx - w//2, y, cx + w//2, y + h
    draw.rounded_rectangle([x1,y1,x2,y2], radius=8, fill=AGENT_BG, outline="#AAAACC", width=2)
    draw.rounded_rectangle([x1+5,y1+5,x2-5,y2-5], radius=5, fill=None, outline="#8899CC", width=1)
    text_center((x1,y1,x2,y2), lines, fb, AGENT_FG)
    return y2

def result_box(cx, y, w, h, lines):
    x1, y1, x2, y2 = cx - w//2, y, cx + w//2, y + h
    draw.rounded_rectangle([x1,y1,x2,y2], radius=6, fill=RESULT_BG, outline=DARK_BLUE, width=2)
    text_center((x1,y1,x2,y2), lines, fb, RESULT_FG)
    return y2

def diamond(cx, y, w, h, lines):
    cy = y + h//2
    pts = [(cx, y), (cx+w//2, cy), (cx, y+h), (cx-w//2, cy)]
    draw.polygon(pts, fill=DECISION_BG, outline=DECISION_BD, width=2)
    text_center((cx-w//2+10, y+10, cx+w//2-10, y+h-10), lines, fn, DECISION_FG)
    return y + h

def api_badge(x, y, text):
    w2, h2 = 148, 28
    draw.rounded_rectangle([x,y,x+w2,y+h2], radius=6, fill=API_BG, outline="#AAAAAA", width=1)
    bb = draw.textbbox((0,0), text, font=fbsm)
    tw = bb[2] - bb[0]
    draw.text((x+(w2-tw)//2, y+(h2-fbsm.size)//2), text, fill=API_FG, font=fbsm)

def arrow_down(cx, y1, y2, color=ARROW_COL):
    draw.line([(cx,y1),(cx,y2)], fill=color, width=2)
    sz = 9
    a = math.pi/2
    draw.polygon([
        (cx - sz*math.cos(a-0.4), y2 - sz*math.sin(a-0.4)),
        (cx, y2),
        (cx - sz*math.cos(a+0.4), y2 - sz*math.sin(a+0.4)),
    ], fill=color)

def arrow_right(x1, y, x2, color=ARROW_COL):
    draw.line([(x1,y),(x2,y)], fill=color, width=2)
    sz = 9
    draw.polygon([
        (x2-sz, y-sz*0.4),
        (x2, y),
        (x2-sz, y+sz*0.4),
    ], fill=color)

def arrow_bent(x1, y1, x2, y2, color=ARROW_COL):
    """L-shaped arrow: right then down, or right then up."""
    mid_x = (x1 + x2) // 2
    draw.line([(x1,y1),(x2,y1)], fill=color, width=2)
    draw.line([(x2,y1),(x2,y2)], fill=color, width=2)
    sz = 9
    if y2 > y1:
        draw.polygon([(x2-sz*0.4, y2-sz), (x2, y2), (x2+sz*0.4, y2-sz)], fill=color)
    else:
        draw.polygon([(x2-sz*0.4, y2+sz), (x2, y2), (x2+sz*0.4, y2+sz)], fill=color)

def dashed_line(x1, y1, x2, y2, color="#888888", dash=8):
    length = math.hypot(x2-x1, y2-y1)
    steps  = int(length / (dash*2))
    for i in range(steps):
        t1 = (i*2*dash) / length
        t2 = min((i*2*dash + dash) / length, 1.0)
        draw.line([
            (x1 + t1*(x2-x1), y1 + t1*(y2-y1)),
            (x1 + t2*(x2-x1), y1 + t2*(y2-y1)),
        ], fill=color, width=2)

# ══════════════════════════════════════════════════════════════════════════
# HEADER BAR
# ══════════════════════════════════════════════════════════════════════════
draw.rectangle([0, 0, W, 58], fill=DARK_BLUE)
title_text = "MedAssist AI — System Architecture & Data Flow"
bb = draw.textbbox((0,0), title_text, font=fxlg)
tw = bb[2] - bb[0]
draw.text(((W-tw)//2, (58-fxlg.size)//2), title_text, fill="white", font=fxlg)

# ══════════════════════════════════════════════════════════════════════════
# COLUMN HEADERS
# ══════════════════════════════════════════════════════════════════════════
# Left column: Patient Flow  |  Right column: Doctor Flow
# Separator line between columns
SEP_X = W // 2
draw.line([(SEP_X, 58), (SEP_X, H-100)], fill=DIVIDER_COL, width=2)

P_CX = 390     # Patient column center
D_CX = 1150    # Doctor column center
BW   = 300     # Standard box width
BH   = 52      # Standard box height
ABH  = 70      # Agent box height

# Column labels
for cx, label in [(P_CX, "PATIENT FLOW"), (D_CX, "DOCTOR FLOW")]:
    draw.rectangle([cx-180, 62, cx+180, 100], fill=MID_BLUE)
    bb = draw.textbbox((0,0), label, font=flg)
    tw = bb[2]-bb[0]
    draw.text((cx-tw//2, 62+(38-flg.size)//2), label, fill="white", font=flg)

# ══════════════════════════════════════════════════════════════════════════
# PATIENT FLOW
# ══════════════════════════════════════════════════════════════════════════
GAP = 18   # vertical gap between boxes (for arrow)
y = 110

# 1. START
y = terminal(P_CX, y, 220, 44, ["Patient Login / Register"]) + GAP
arrow_down(P_CX, y-GAP, y)

# 2. Profile Setup
y = process_box(P_CX, y, BW, BH,
    ["Patient Profile Setup",
     "(Age, Weight, Height, Blood Group,",
     "Allergies, Current Medications)"],
    font=fsm) + GAP
arrow_down(P_CX, y-GAP, y)

# 3. Symptom Intake
y = process_box(P_CX, y, BW, BH,
    ["Symptom Intake Wizard",
     "(7 Body Systems · 36 Symptoms",
     "Duration · Severity · Onset)"],
    font=fsm) + GAP
arrow_down(P_CX, y-GAP, y)

# 4. Diagnostic Agent
diag_y = y
y = agent_box(P_CX, y, BW, ABH,
    ["\u2699  Diagnostic Agent",
     "LLM Tool-Use Loop"]) + GAP
# API badge (right side)
api_badge(P_CX + BW//2 + 14, diag_y + 22, "ICD-10 API (NIH)")
draw.line([(P_CX+BW//2, diag_y+22+14), (P_CX+BW//2+14, diag_y+22+14)], fill="#888888", width=1)

arrow_down(P_CX, y-GAP, y)

# 5. Disease Results
y = result_box(P_CX, y, BW, BH,
    ["Top 5 Predicted Diseases",
     "ICD-10 Codes · Probability Scores"]) + GAP
arrow_down(P_CX, y-GAP, y)

# 6. Select Disease + Blood Tests
y = process_box(P_CX, y, BW, BH,
    ["Patient Selects Disease",
     "Recommended Blood Test List"],
    font=fsm) + GAP
arrow_down(P_CX, y-GAP, y)

# 7. Upload Blood Report
y = process_box(P_CX, y, BW, BH,
    ["Upload Blood Report",
     "(PDF or Image Scan)"],
    font=fsm) + GAP
arrow_down(P_CX, y-GAP, y)

# 8. OCR
y = process_box(P_CX, y, BW, BH,
    ["OCR: Extract Lab Values",
     "(pdf-parse + Groq Vision)"],
    font=fsm) + GAP
arrow_down(P_CX, y-GAP, y)

# 9. Blood Report Agent
br_y = y
y = agent_box(P_CX, y, BW, ABH,
    ["\u2699  Blood Report Agent",
     "LLM Tool-Use Loop"]) + GAP
# API badges
api_badge(P_CX + BW//2 + 14, br_y + 8,  "OpenFDA API")
api_badge(P_CX + BW//2 + 14, br_y + 42, "RxNorm API")
draw.line([(P_CX+BW//2, br_y+22), (P_CX+BW//2+14, br_y+22)], fill="#888888", width=1)
draw.line([(P_CX+BW//2, br_y+56), (P_CX+BW//2+14, br_y+56)], fill="#888888", width=1)

arrow_down(P_CX, y-GAP, y)

# 10. Medication plan
y = result_box(P_CX, y, BW, BH,
    ["Medication Plan +",
     "Complexity Score"]) + GAP
arrow_down(P_CX, y-GAP, y)

# 11. Decision diamond
diamond_top = y
dia_h = 68
y = diamond(P_CX, y, 260, dia_h,
    ["High Complexity", "or Critical Interaction?"])

# NO branch → down to END
no_y = y + GAP
arrow_down(P_CX, y, no_y)
# Label "No"
draw.text((P_CX+6, y+4), "No", fill=DECISION_FG, font=fbsm)

y = no_y
y = terminal(P_CX, y, 200, 40, ["END — Patient Flow"]) + 0

# YES branch → right to Doctor Finder
yes_x = P_CX + 130          # right tip of diamond
finder_cx = P_CX + 130 + 190
finder_y  = diamond_top + dia_h//2 - BH//2
# draw horizontal + down arrow
draw.line([(yes_x, diamond_top+dia_h//2),
           (finder_cx-BW//2, diamond_top+dia_h//2)], fill=ARROW_COL, width=2)
draw.polygon([
    (finder_cx-BW//2, diamond_top+dia_h//2 - 6),
    (finder_cx-BW//2, diamond_top+dia_h//2 + 6),
    (finder_cx-BW//2+9, diamond_top+dia_h//2),
], fill=ARROW_COL)
draw.text((yes_x+4, diamond_top+dia_h//2-16), "Yes", fill=DECISION_FG, font=fbsm)
# Doctor Finder box
f_x1 = finder_cx - BW//2
f_y1 = finder_y
f_x2 = finder_cx + BW//2
f_y2 = finder_y + BH + 10
draw.rounded_rectangle([f_x1, f_y1, f_x2, f_y2], radius=6,
    fill="#16537E", outline=MID_BLUE, width=2)
text_center((f_x1, f_y1, f_x2, f_y2),
    ["\U0001F5FA  Doctor Finder Map",
     "(OpenStreetMap · Nominatim)"], fbsm, "white")

# ══════════════════════════════════════════════════════════════════════════
# DOCTOR FLOW
# ══════════════════════════════════════════════════════════════════════════
y = 110

y = terminal(D_CX, y, 220, 44, ["Doctor Login / Register"]) + GAP
arrow_down(D_CX, y-GAP, y)

y = process_box(D_CX, y, BW, BH,
    ["Doctor Dashboard"],
    font=fb) + GAP
arrow_down(D_CX, y-GAP, y)

y = process_box(D_CX, y, BW, BH,
    ["Enter Patient Summary",
     "+ Prescribed Blood Tests"],
    font=fsm) + GAP
arrow_down(D_CX, y-GAP, y)

da_y = y
y = agent_box(D_CX, y, BW, ABH,
    ["\u2699  Doctor Assist Agent",
     "LLM Tool-Use Loop"]) + GAP
api_badge(D_CX + BW//2 + 14, da_y + 22, "ICD-10 API (NIH)")
draw.line([(D_CX+BW//2, da_y+22+14), (D_CX+BW//2+14, da_y+22+14)], fill="#888888", width=1)

arrow_down(D_CX, y-GAP, y)

y = result_box(D_CX, y, BW, BH,
    ["Missing Blood Tests",
     "Urgency: Routine / Urgent / Critical"]) + GAP
arrow_down(D_CX, y-GAP, y)

y = process_box(D_CX, y, BW, BH,
    ["Doctor Reviews AI Suggestions",
     "+ Audit Trail Available"],
    font=fsm) + GAP
arrow_down(D_CX, y-GAP, y)

terminal(D_CX, y, 200, 40, ["END — Doctor Flow"])

# ══════════════════════════════════════════════════════════════════════════
# SHARED INFRASTRUCTURE BANNER (bottom)
# ══════════════════════════════════════════════════════════════════════════
SHARED_Y = H - 96
draw.rectangle([0, SHARED_Y, W, H], fill="#EEF3FB")
draw.line([(0, SHARED_Y), (W, SHARED_Y)], fill=MID_BLUE, width=2)

# Label
draw.text((20, SHARED_Y + 8), "SHARED INFRASTRUCTURE:", fill=DARK_BLUE, font=flg)

infra_items = [
    ("SSE Agent Status Tracker", "Real-time step-by-step UI"),
    ("PostgreSQL Agent Audit Log", "Full tool call history"),
    ("JWT Auth Middleware", "Role-based access"),
    ("React + Vite Frontend", "Express.js Backend"),
]
ix = 280
for title, sub in infra_items:
    iw = 270
    draw.rounded_rectangle([ix, SHARED_Y+6, ix+iw, H-8],
        radius=8, fill=MID_BLUE, outline=DARK_BLUE, width=1)
    bb = draw.textbbox((0,0), title, font=fbsm)
    tw = bb[2]-bb[0]
    draw.text((ix+(iw-tw)//2, SHARED_Y+14), title, fill="white", font=fbsm)
    bb2 = draw.textbbox((0,0), sub, font=fsm)
    tw2 = bb2[2]-bb2[0]
    draw.text((ix+(iw-tw2)//2, SHARED_Y+38), sub, fill="#D6E4F7", font=fsm)
    ix += iw + 14

# ══════════════════════════════════════════════════════════════════════════
# LEGEND
# ══════════════════════════════════════════════════════════════════════════
LEG_Y = SHARED_Y - 46
draw.rectangle([0, LEG_Y, W, SHARED_Y-2], fill="#F5F5F5")
draw.line([(0, LEG_Y), (W, LEG_Y)], fill="#CCCCCC", width=1)

legend = [
    (AGENT_BG,    AGENT_FG,    " AI Agent   "),
    (RESULT_BG,   RESULT_FG,   " Results    "),
    (PROCESS_BG,  PROCESS_FG,  " Process    "),
    (TERMINAL_BG, TERMINAL_FG, " Start / End"),
    (DECISION_BG, DECISION_FG, " Decision   "),
    (API_BG,      API_FG,      " External API"),
]
lx = 20
for bg, fg, label in legend:
    bw2, bh2 = 90, 26
    draw.rounded_rectangle([lx, LEG_Y+8, lx+bw2, LEG_Y+8+bh2],
        radius=5, fill=bg, outline="#888888", width=1)
    bb = draw.textbbox((0,0), label.strip(), font=fsm)
    tw = bb[2]-bb[0]
    draw.text((lx+(bw2-tw)//2, LEG_Y+8+(bh2-fsm.size)//2),
              label.strip(), fill=fg, font=fsm)
    lx += bw2 + 16

# ══════════════════════════════════════════════════════════════════════════
# Save flowchart image
# ══════════════════════════════════════════════════════════════════════════
out_img = r"C:\prsnl_doc\CS595\Project\flowchart.png"
img.save(out_img, dpi=(150, 150))
print(f"Flowchart saved: {out_img}")

# ══════════════════════════════════════════════════════════════════════════
# Insert into the existing docx
# ══════════════════════════════════════════════════════════════════════════
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DOCX_PATH = r"C:\prsnl_doc\CS595\Project\MedAssist_AI_Project_Proposal_Updated.docx"
doc = Document(DOCX_PATH)

DARK_BLUE_C = RGBColor(0x1F, 0x49, 0x7D)
MID_BLUE_C  = RGBColor(0x2E, 0x74, 0xB5)

def add_heading1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1F497D')
    pBdr.append(bottom)
    pPr.append(pBdr)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = DARK_BLUE_C
    run.font.name = "Calibri"
    return p

# Add page break
doc.add_page_break()

# Section heading
add_heading1(doc, "3. System Architecture & Data Flow Diagram")

# Caption paragraph
cap = doc.add_paragraph()
cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
cap.paragraph_format.space_before = Pt(6)
cap.paragraph_format.space_after  = Pt(6)
r = cap.add_run(
    "Figure 1 — End-to-end data flow for the Patient Diagnostic Pipeline (left) "
    "and Doctor Assist Pipeline (right), including AI agents, external API calls, "
    "and shared infrastructure."
)
r.italic = True
r.font.size = Pt(10)
r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
r.font.name = "Calibri"

# Insert image — centered, scaled to fit page width
pic_para = doc.add_paragraph()
pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
pic_para.paragraph_format.space_after = Pt(8)
run = pic_para.add_run()
# Page content width ≈ 6.5 inches (8.5" - 2" margins)
run.add_picture(out_img, width=Inches(6.5))

# Brief description
desc = doc.add_paragraph()
desc.paragraph_format.space_before = Pt(4)
desc.paragraph_format.space_after  = Pt(6)
r = desc.add_run(
    "The diagram above illustrates the complete request lifecycle for both user roles. "
    "The Patient Flow (left) begins with symptom collection and progresses through the "
    "Diagnostic Agent — which calls the NIH ICD-10 API to verify disease codes — to the "
    "Blood Report Agent, which invokes OpenFDA and RxNorm APIs to validate medications and "
    "check drug interactions before delivering a personalized analysis. High-complexity cases "
    "automatically trigger the geolocation-based Doctor Finder. The Doctor Flow (right) routes "
    "physician input through the Doctor Assist Agent to detect missing blood tests using the same "
    "ICD-10 grounding. All three agents stream live progress events to the frontend via "
    "Server-Sent Events (SSE) and log every tool call to the PostgreSQL audit table."
)
r.font.size = Pt(11)
r.font.name = "Calibri"

doc.save(DOCX_PATH)
print(f"Document updated: {DOCX_PATH}")
