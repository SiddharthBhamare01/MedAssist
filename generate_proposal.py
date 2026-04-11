from docx import Document
from docx.shared import Pt, RGBColor, Inches, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Inches(8.5)
section.page_height = Inches(11)
section.left_margin   = Inches(1)
section.right_margin  = Inches(1)
section.top_margin    = Inches(1)
section.bottom_margin = Inches(1)

# ── Style helpers ─────────────────────────────────────────────────────────────
DARK_BLUE = RGBColor(0x1F, 0x49, 0x7D)
MID_BLUE  = RGBColor(0x2E, 0x74, 0xB5)
BLACK     = RGBColor(0x00, 0x00, 0x00)

def set_run(run, bold=False, italic=False, size=11, color=BLACK, font="Calibri"):
    run.bold   = bold
    run.italic = italic
    run.font.size  = Pt(size)
    run.font.color.rgb = color
    run.font.name  = font

def heading1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    # bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1F497D')
    pBdr.append(bottom)
    pPr.append(pBdr)
    run = p.add_run(text)
    set_run(run, bold=True, size=14, color=DARK_BLUE)
    return p

def heading2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    set_run(run, bold=True, size=12, color=MID_BLUE)
    return p

def heading3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    set_run(run, bold=True, size=11, color=BLACK)
    return p

def body(text, bold_prefix=None):
    """Add a body paragraph; if bold_prefix given, render it bold then rest normal."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        set_run(r, bold=True, size=11)
        r2 = p.add_run(text)
        set_run(r2, size=11)
    else:
        r = p.add_run(text)
        set_run(r, size=11)
    return p

def bullet(text, bold_prefix=None, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent   = Inches(0.25 + level * 0.25)
    p.paragraph_format.space_before  = Pt(1)
    p.paragraph_format.space_after   = Pt(2)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        set_run(r, bold=True, size=11)
        r2 = p.add_run(text)
        set_run(r2, size=11)
    else:
        r = p.add_run(text)
        set_run(r, size=11)
    return p

def checkbox(checked, text):
    sym = "\u2611" if checked else "\u2610"
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    r = p.add_run(f"{sym}  {text}")
    set_run(r, size=11)
    return p

def page_break():
    doc.add_page_break()

def divider():
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'AAAAAA')
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)

# ══════════════════════════════════════════════════════════════════════════════
# TITLE BLOCK
# ══════════════════════════════════════════════════════════════════════════════
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_p.paragraph_format.space_before = Pt(0)
title_p.paragraph_format.space_after  = Pt(6)
r = title_p.add_run("MedAssist AI")
set_run(r, bold=True, size=22, color=DARK_BLUE)

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_p.paragraph_format.space_after = Pt(2)
r = sub_p.add_run("Intelligent Medical Diagnostic & Prescription Support Web Application")
set_run(r, bold=False, size=13, color=MID_BLUE)

meta_p = doc.add_paragraph()
meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta_p.paragraph_format.space_after = Pt(12)
r = meta_p.add_run("CS 595: Medical Informatics and AI  |  Team 2  |  Project Proposal")
set_run(r, italic=True, size=10, color=RGBColor(0x60, 0x60, 0x60))

divider()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — PROJECT DEFINITION
# ══════════════════════════════════════════════════════════════════════════════
heading1("1. Project Definition")

heading2("a. Project Title")
body("Project ID: Team 2", bold_prefix="")
p = doc.add_paragraph()
r = p.add_run("Project Name: ")
set_run(r, bold=True, size=11)
r2 = p.add_run("MedAssist AI — Intelligent Medical Diagnostic & Prescription Support Web Application")
set_run(r2, size=11)
p.paragraph_format.space_after = Pt(4)

heading2("b. Author(s)")
authors = [
    "Siddharth Bhamare (A20582786)",
    "Jivan Singh (A20581464)",
    "Gayatri Sanjay Gaikwad (A20598464)",
    "Vaishnav Bhujbal (A20579636)",
]
for i, a in enumerate(authors, 1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    r = p.add_run(f"{i}.  {a}")
    set_run(r, size=11)

heading2("c. LOF Pillar")
checkbox(True,  "Patient Engagement")
checkbox(False, "Data Analysis & Population Health")
checkbox(False, "Medical Education")

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2 — INTRODUCTION
# ══════════════════════════════════════════════════════════════════════════════
heading1("2. Introduction")

heading2("a. Problem Description")
body(
    "Patients without immediate access to medical professionals often struggle to interpret their "
    "symptoms, understand which diagnostic tests are relevant to their condition, and make sense of "
    "blood report results. This leads to delayed care, uninformed self-treatment, and avoidable "
    "complications. On the clinical side, physicians working under time pressure may occasionally "
    "omit relevant blood tests from a workup, resulting in incomplete or delayed diagnoses."
)
body(
    "The opportunity this project addresses is the gap between raw patient data — symptoms and lab "
    "reports — and actionable, personalized medical guidance. MedAssist AI bridges this gap by "
    "deploying an AI-powered web application that walks patients through a structured, intelligent "
    "diagnostic pipeline: symptom collection → AI disease prediction → blood test guidance → blood "
    "report analysis → medication recommendations → doctor referral when needed. A secondary module "
    "serves physicians by reviewing their prescriptions and flagging potentially missing blood tests."
)
body(
    "The work performed includes full-stack web application development using a React frontend, "
    "Node.js/Express backend, PostgreSQL database, and Groq-hosted LLM agents (Llama 3.3-70B) that "
    "autonomously call real medical APIs (OpenFDA, RxNorm, NIH ICD-10) to ground all recommendations "
    "in verified, FDA-approved clinical data. The solution takes the form of a role-based web "
    "application accessible via any modern browser."
)

heading2("b. Purpose")
body(
    "MedAssist AI is an intelligent, multi-role web application that empowers patients to understand "
    "their health through AI-guided symptom analysis and blood report interpretation, while providing "
    "doctors with an AI-powered tool to verify and validate their prescribed diagnostic tests against "
    "clinical standards."
)
heading3("Goals and Objectives:")
bullet("Provide patients with a step-by-step AI diagnostic experience from symptom input to personalized, FDA-grounded medication recommendations.")
bullet("Use agentic AI (LLM with Tool Use) to autonomously verify disease codes, validate lab values, and check drug interactions in real time before surfacing any clinical suggestion.")
bullet("Enable physicians to leverage AI-assisted clinical decision support to reduce diagnostic gaps in blood test prescriptions.")
bullet("Demonstrate a production-quality, architecturally significant AI system appropriate for academic evaluation in CS 595: Medical Informatics and AI.")

heading2("c. Scope")
heading3("Included in Scope:")
scope_in = [
    "User registration and authentication for two roles: Patient and Doctor (JWT-based).",
    "Patient symptom collection wizard (multi-step form covering 7 body systems: General, Respiratory, Digestive, Neurological, Musculoskeletal, Skin, Urinary).",
    "AI-powered disease prediction returning top 5 candidates with ICD-10 codes and probability scores.",
    "Blood test recommendation per selected disease with clinical rationale.",
    "Blood report image/PDF upload with OCR extraction of lab values (pdf-parse + Groq).",
    "AI blood report analysis with personalized, weight/height-adjusted medication recommendations.",
    "Real-time FDA drug validation via OpenFDA API and drug interaction checks via RxNorm API.",
    "Complexity scoring and automated doctor referral trigger for high-risk cases.",
    "Geolocation-based doctor finder using OpenStreetMap / Nominatim (free, no API key required).",
    "Doctor-facing AI tool to detect missing blood tests from an existing prescription.",
    "Real-time agent step tracker UI showing live tool call progress (via Server-Sent Events).",
    "Agent audit log stored in the database for academic transparency and professor review.",
    "Deployment via Docker Compose (single machine) or Vercel + Railway.",
    "Disclaimer banner identifying the app as an educational tool, not a substitute for professional medical advice.",
]
for item in scope_in:
    bullet(item)

heading3("Excluded from Scope:")
scope_out = [
    "Real Electronic Health Record (EHR) integration or FHIR/HL7 connectivity.",
    "HIPAA compliance or handling of real protected health information (PHI).",
    "Native mobile applications (iOS / Android).",
    "Telemedicine or real-time doctor–patient chat or video.",
    "Prescription writing or e-prescribing functionality.",
    "Integration with pharmacy or insurance systems.",
    "Training or fine-tuning of custom machine learning models.",
    "Radiology imaging analysis (X-ray, MRI, CT scan).",
    "Multi-language / internationalization support.",
]
for item in scope_out:
    bullet(item)

heading2("d. Audience")
heading3("Primary Users — Patients")
bullet("Who: Adults (18+) seeking preliminary guidance on symptoms or help interpreting blood test results.")
bullet("Health Literacy: Varying; the application must present information in plain, jargon-free language.")
bullet("Access: Mobile and desktop browser; no app installation required.")
bullet("Motivation: Driven by a specific health concern; may be anxious or overwhelmed by medical terminology.")
bullet("Technical Proficiency: Assumed to be general internet users, not technically specialized.")
bullet("Key Need: Clear, step-by-step guidance; personalized output based on their age, weight, and medical history; confidence that suggestions are clinically grounded.")

heading3("Secondary Users — Physicians / Clinicians")
bullet("Who: General practitioners or specialist doctors seeking AI-assisted support for clinical decision-making.")
bullet("Health Literacy: High; fluent in diagnostic terminology and blood test nomenclature.")
bullet("Access: Desktop browser in a clinical or office setting.")
bullet("Motivation: Time-constrained; want fast, accurate suggestions with traceable clinical reasoning.")
bullet("Key Need: Minimal UI friction; actionable, prioritized output; audit trail for AI-generated suggestions.")

heading3("Indirect Stakeholder")
bullet("Leap Of Faith (LOF) — evaluating the system's architecture, AI integration quality, and academic rigor.")

heading2("e. Overview")
body(
    "MedAssist AI is a full-stack web application organized around two user roles: Patient and Doctor. "
    "Each role has a dedicated flow with purpose-built pages and AI agents that power the core functionality."
)
body(
    "Patient Flow: Patients begin with a multi-step intake wizard that collects personal health "
    "information (age, weight, height, blood group, existing conditions, allergies, current medications) "
    "and current symptoms across seven body systems. This data is passed to the Diagnostic Agent — an "
    "autonomous AI agent built on a Groq-hosted LLM with Tool Use — which reasons over the symptoms, "
    "calls the NIH ICD-10 API to verify official disease codes, and returns the top five predicted "
    "diseases ranked by probability score."
)
body(
    "The patient selects a disease to investigate further and receives a clinically relevant list of "
    "recommended blood tests. They then upload a photo or scan of their blood report. OCR (pdf-parse + "
    "Groq) extracts all lab values, and the Blood Report Agent validates each value against "
    "demographic-adjusted clinical reference ranges, calls the OpenFDA API to identify FDA-approved "
    "treatments, checks drug interactions via RxNorm against the patient's existing medications, and "
    "produces a fully personalized analysis including weight-adjusted medication recommendations and a "
    "complexity score. When complexity is high or critical drug interactions are detected, the system "
    "automatically triggers a doctor referral prompt backed by a geolocation-based doctor finder "
    "powered by OpenStreetMap."
)
body(
    "Doctor Flow: Physicians access a separate dashboard where they enter a patient summary and their "
    "current list of prescribed blood tests. The Doctor Assist Agent uses the same tool-use architecture "
    "to identify the clinical context via ICD-10 lookup, enumerate the standard-of-care test panel for "
    "that condition, compare against the existing prescription, and return a prioritized list of missing "
    "tests with urgency levels (Routine / Urgent / Critical)."
)
body(
    "All three agents emit real-time progress events to the frontend via Server-Sent Events (SSE), "
    "rendered in a live Agent Status Tracker component so users can observe each reasoning step as it "
    "executes. Every agent run — including each tool called, its inputs and outputs, and the number of "
    "reasoning loops — is logged to the database in an audit table, providing full transparency for "
    "academic review."
)

# ══════════════════════════════════════════════════════════════════════════════
# PAGE BREAK BEFORE ADDENDUM
# ══════════════════════════════════════════════════════════════════════════════
page_break()

# ══════════════════════════════════════════════════════════════════════════════
# ADDENDUM HEADER
# ══════════════════════════════════════════════════════════════════════════════
add_title = doc.add_paragraph()
add_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_title.paragraph_format.space_before = Pt(0)
add_title.paragraph_format.space_after  = Pt(4)
r = add_title.add_run("CS 595 Project Proposal — Addendum")
set_run(r, bold=True, size=16, color=DARK_BLUE)

add_sub = doc.add_paragraph()
add_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_sub.paragraph_format.space_after = Pt(12)
r = add_sub.add_run("Pillar Alignment and Sprint Execution Plan")
set_run(r, italic=True, size=12, color=MID_BLUE)

divider()

# ══════════════════════════════════════════════════════════════════════════════
# ADDENDUM SECTION 1 — PRIMARY PILLAR ALIGNMENT
# ══════════════════════════════════════════════════════════════════════════════
heading1("1. Primary Pillar Alignment")

body("Selected Pillar:")
checkbox(True,  "Patient Engagement")
checkbox(False, "Data Analytics / Population Health")
checkbox(False, "Medical Education")

# ══════════════════════════════════════════════════════════════════════════════
# ADDENDUM SECTION 2 — LOF HEALTHCARE PILLAR JUSTIFICATION
# ══════════════════════════════════════════════════════════════════════════════
heading1("2. LOF Healthcare Pillar Justification")

body(
    "MedAssist AI belongs under the Patient Engagement pillar because its primary user is the patient "
    "— an adult seeking to understand their own health without immediate access to a physician. The "
    "core healthcare problem the project addresses is the disconnect between raw health data (symptoms "
    "and lab results) and the patient's ability to act on that data confidently and safely."
)
body(
    "The project creates value within the Patient Engagement pillar by actively involving patients in "
    "every step of their diagnostic journey rather than passively delivering a result. Through a "
    "structured, multi-step intake wizard and an AI-guided analysis pipeline, the application builds "
    "health literacy: patients learn which tests relate to their symptoms, what their lab values mean "
    "relative to their own demographics, and when the complexity of their case warrants seeing a doctor."
)
body(
    "The AI agents — powered by a Groq-hosted LLM with Tool Use calling FDA, RxNorm, and ICD-10 APIs "
    "— ensure that every recommendation surfaced to the patient is grounded in clinically verified, "
    "FDA-approved data, directly reinforcing patient trust and informed decision-making. The secondary "
    "physician module further supports engagement by helping doctors deliver more complete diagnostic "
    "workups to their patients, closing gaps that would otherwise delay care."
)

# ══════════════════════════════════════════════════════════════════════════════
# ADDENDUM SECTION 3 — EVIDENCE FROM PROPOSAL
# ══════════════════════════════════════════════════════════════════════════════
heading1("3. Evidence from the Proposal")

heading2("Evidence 1 — Target Users")
body(
    "The proposal explicitly identifies adults (18+) with varying health literacy as the primary "
    "audience, and specifies that the application must present information in plain, jargon-free "
    "language with step-by-step guidance. This audience definition is the hallmark of a Patient "
    "Engagement system: the design priority is empowering the patient, not aggregating population "
    "data or training clinicians."
)

heading2("Evidence 2 — Patient-Centric Feature Set")
body(
    "The full patient workflow — symptom intake wizard → AI disease prediction with ICD-10 codes → "
    "blood test recommendations → blood report upload and OCR → personalized, weight-adjusted "
    "medication recommendations → complexity score → doctor referral trigger — is entirely oriented "
    "around guiding a single patient through their own health question. Every feature exists to "
    "improve individual patient understanding and action, not to analyze trends across a population."
)

heading2("Evidence 3 — Real-Time Agent Status Tracker and Audit Log")
body(
    "The live Agent Status Tracker (SSE-based) shows patients each reasoning step the AI takes on "
    "their behalf, including which medical APIs were called and what data was retrieved. This "
    "transparency feature is a deliberate Patient Engagement design choice: it builds trust, reduces "
    "anxiety, and educates the patient about how the AI arrived at its suggestions — directly "
    "supporting informed patient participation in their own care."
)

# ══════════════════════════════════════════════════════════════════════════════
# ADDENDUM SECTION 4 — SPRINT EXECUTION PLAN
# ══════════════════════════════════════════════════════════════════════════════
heading1("4. Sprint Execution Plan")

heading2("Sprint 1 — Core Patient Diagnostic Pipeline")
body(
    "Sprint 1 focuses on delivering the end-to-end patient flow from symptom input to personalized "
    "medication recommendations. The following capabilities will be implemented and integrated:"
)
sprint1_items = [
    "User registration and JWT authentication for Patient and Doctor roles.",
    "Patient profile setup and multi-step symptom intake wizard (7 body systems, 36 symptoms with duration, severity, and onset).",
    "Diagnostic Agent: LLM tool-use loop calling NIH ICD-10 API, returning top 5 predicted diseases with probability scores and ICD codes.",
    "Disease results page with disease selection and recommended blood test list.",
    "Blood report upload UI supporting PDF and image formats; OCR extraction of lab values via pdf-parse and Groq.",
    "Blood Report Agent: lab value validation against demographic-adjusted reference ranges, OpenFDA drug lookup, RxNorm drug interaction check, weight-adjusted medication recommendations, and complexity scoring.",
    "Blood report analysis results page with full medication plan and referral trigger for high-complexity cases.",
    "Real-time Agent Status Tracker UI (SSE) showing live tool call progress for both agents.",
    "Agent audit log stored in PostgreSQL for all agent runs.",
]
for item in sprint1_items:
    bullet(item)

heading2("Sprint 2 — Doctor Module, Doctor Finder, and Polish")
body(
    "Sprint 2 expands the application with the physician-facing module, geolocation features, and "
    "production-readiness improvements:"
)
sprint2_items = [
    "Doctor dashboard: patient summary and prescribed blood test input form.",
    "Doctor Assist Agent: ICD-10 context lookup, standard-of-care test panel comparison, missing test identification with urgency levels (Routine / Urgent / Critical).",
    "Geolocation-based doctor finder using OpenStreetMap / Nominatim and Leaflet.js.",
    "Agent log viewer page for doctors and patients to review full AI reasoning history.",
    "UI polish: responsive layout, mobile-friendly design, accessibility improvements, disclaimer banner.",
    "End-to-end integration testing across the full patient and doctor flows.",
    "Documentation: AI prompt documentation (docs/ai-prompts.md) and project info template.",
    "Demo preparation and final submission packaging.",
]
for item in sprint2_items:
    bullet(item)

# ══════════════════════════════════════════════════════════════════════════════
# ADDENDUM SECTION 5 — SCOPE CONFIRMATION
# ══════════════════════════════════════════════════════════════════════════════
heading1("5. Scope Confirmation")

checkbox(True,  "This sprint plan is fully within our original scope.")
checkbox(False, "This sprint plan represents a refinement or narrowing of our original scope.")

body(
    "All features listed in Sprint 1 and Sprint 2 correspond directly to items enumerated in the "
    "Included in Scope section of the original proposal. No new features have been added, and no "
    "originally scoped items have been removed. The sprint breakdown simply sequences the delivery "
    "of the full original scope across two development cycles."
)

# ══════════════════════════════════════════════════════════════════════════════
# ADDENDUM SECTION 6 — HEALTHCARE DATA, STANDARDS, OR CLINICAL CONTEXT
# ══════════════════════════════════════════════════════════════════════════════
heading1("6. Healthcare Data, Standards, and Clinical Context")

body(
    "MedAssist AI integrates four established healthcare data standards and free public clinical APIs "
    "to ground all AI-generated recommendations in verified medical knowledge:"
)

heading2("ICD-10 (International Classification of Diseases, 10th Revision)")
body(
    "Used by the Diagnostic Agent to map predicted diseases to official ICD-10 codes via the NIH "
    "ClinicalTables API. Every disease returned to the patient carries a validated ICD-10 code, "
    "ensuring the diagnostic output aligns with the international standard for disease classification "
    "used in clinical practice worldwide."
)

heading2("RxNorm (National Library of Medicine)")
body(
    "Used by the Blood Report Agent to check drug-drug interactions between AI-recommended medications "
    "and the patient's existing medication list. RxNorm is the NLM's standard for clinical drug "
    "nomenclature and is the authoritative source for drug interaction data in U.S. clinical systems."
)

heading2("OpenFDA (U.S. Food and Drug Administration Drug Database)")
body(
    "Used by the Blood Report Agent to validate that all recommended treatments are FDA-approved for "
    "the identified condition. The OpenFDA API provides access to the official FDA drug label database, "
    "ensuring no unapproved or off-label medications are surfaced without explicit flagging."
)

heading2("Demographic-Adjusted Clinical Reference Ranges")
body(
    "The Blood Report Agent applies age-, sex-, and weight-adjusted reference ranges for all common "
    "blood panel values (CBC, metabolic panel, lipid panel, thyroid, etc.). These ranges are derived "
    "from published clinical laboratory guidelines and are used to classify each patient lab value as "
    "normal, borderline, or abnormal in the context of that patient's specific demographics — "
    "reflecting real clinical laboratory interpretation practice."
)

heading2("Clinical Workflow Alignment")
body(
    "Even where a formal interoperability standard such as HL7 FHIR is not used, the application "
    "mirrors real clinical workflows: symptom triage → differential diagnosis → laboratory workup → "
    "result interpretation → treatment planning → specialist referral. The Doctor Assist Agent "
    "mirrors the clinical practice of a physician reviewing their own order set against evidence-based "
    "test panels for a given diagnosis — a standard quality-assurance step in clinical decision support."
)

# ══════════════════════════════════════════════════════════════════════════════
# FOOTER NOTE
# ══════════════════════════════════════════════════════════════════════════════
divider()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(
    "MedAssist AI is an educational prototype. It is not a substitute for professional medical advice, "
    "diagnosis, or treatment. All AI-generated suggestions should be reviewed by a qualified healthcare provider."
)
set_run(r, italic=True, size=9, color=RGBColor(0x80, 0x80, 0x80))

# ══════════════════════════════════════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════════════════════════════════════
out = r"C:\prsnl_doc\CS595\Project\MedAssist_AI_Project_Proposal_Updated.docx"
doc.save(out)
print(f"Saved: {out}")
